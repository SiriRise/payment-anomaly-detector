import os
import re
import sys
import datetime
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")


def _normalize_val(val):
    if val is None:
        return None
    s = str(val).strip()
    return s if s else None


def _normalize_for_compare(text):
    if not text:
        return ""

    ocr_fix = str.maketrans(
        "HBCOPEKMTXAEhbcopkmtxae",
        "НВСОРЕКМТХАЕнвсоркмтхае"
    )
    text = text.translate(ocr_fix)

    text = re.sub(r'[^а-яёa-z\s]', '', text.lower())
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _names_match(h_val, qr_val, min_overlap=0.5):
    norm_h = _normalize_for_compare(h_val)
    norm_qr = _normalize_for_compare(qr_val)
    if not norm_h or not norm_qr:
        return False
    if norm_h == norm_qr:
        return True
    tokens_h = norm_h.split()
    tokens_qr = norm_qr.split()
    matched = 0
    for t1 in tokens_h:
        for t2 in tokens_qr:
            if t1 == t2 or t1.startswith(t2) or t2.startswith(t1):
                matched += 1
                break

    overlap = matched / max(len(tokens_h), len(tokens_qr))
    return overlap >= min_overlap


def parse_number(text):
    if not text:
        return None
    try:
        cleaned = str(text).strip().replace(' ', '').replace(',', '.')
        cleaned = re.sub(r'[^\d.\-]', '', cleaned)
        if not cleaned or cleaned == '-' or cleaned == '.':
            return None
        return float(cleaned)
    except:
        return None


def get_val(row_cells, col_num):
    for cell in row_cells:
        if cell.get('col_num') == col_num:
            txt = cell.get('text', '').strip()
            if not txt or txt == '-':
                return None
            return parse_number(txt)
    return None


def check_table_formulas(tables: list) -> dict:
    errors = []
    stats = {
        "tables_checked": 0,
        "rows_checked": 0,
        "errors_count": 0
    }
    for table in tables:
        page_num = table.get('page', '?')
        table_idx = table.get('table_idx', '?')
        cells = table.get('cells_data', [])

        if not cells:
            continue
        rows = {}
        for cell in cells:
            r = cell.get('row_num')
            if r is not None:
                rows.setdefault(r, []).append(cell)

        if not rows:
            continue

        stats["tables_checked"] += 1

        for r_num, row_cells in sorted(rows.items()):

            c1 = get_val(row_cells, 1)
            c2 = get_val(row_cells, 2)
            c3 = get_val(row_cells, 3)  # Тариф
            c4 = get_val(row_cells, 4)
            c5 = get_val(row_cells, 5)  # Объем
            c6 = get_val(row_cells, 6)  # Коэфф
            c7 = get_val(row_cells, 7)  # Начислено
            c8 = get_val(row_cells, 8)  # Пени/доп
            c9 = get_val(row_cells, 9)  # Льготы
            c10 = get_val(row_cells, 10)  # Итого
            c11 = get_val(row_cells, 11)  # Пред. начислено
            c12 = get_val(row_cells, 12)  # Пред. оплачено
            c13 = get_val(row_cells, 13)  # Долг/Аванс
            c14 = get_val(row_cells, 14)  # Всего к оплате

            max_col = max((c.get('col_num', 0) for c in row_cells), default=0)

            if max_col == 15:
                #print(c1,c2,c3,c4,c5,c6,c7,c8,c9,c10,c11,c12,c13,c14)
                #print(table_idx,r_num,max_col)

                row_errors = []

                if c7 is not None and c3 is not None and c5 is not None:
                    calc = (c3 or 0) * (c5 or 0) * (c6 or 1)
                    if abs(calc - (c7 or 0)) > 0.50:
                        row_errors.append({
                            "formula": "C7 = C3 * C5 * C6",
                            "expected": calc,
                            "actual": (c7 or 0),
                            "diff": abs(calc - (c7 or 0))
                        })
                else:
                    print(table_idx, r_num, " НЕ ПРОВКРИЛ")
                if c7 is not None and c10 is not None:
                    calc = (c7 or 0) + (c8 or 0) - (c9 or 0)
                    if abs(calc - c10) > 0.50:
                        row_errors.append({
                            "formula": "C10 = C7 + C8 - C9",
                            "expected": calc,
                            "actual": c10,
                            "diff": abs(calc - c10)
                        })
                else:
                    print(table_idx, r_num, " НЕ ПРОВКРИЛ")
                if c11 is not None and c12 is not None and c13 is not None:
                    calc = (c11 or 0) - (c12 or 0)
                    if abs(calc - c13) > 0.50:
                        row_errors.append({
                            "formula": "C13 = C11 - C12",
                            "expected": calc,
                            "actual": c13,
                            "diff": abs(calc - c13)
                        })
                else:
                    print(table_idx, r_num, " НЕ ПРОВКРИЛ")
                if c10 is not None and c14 is not None:
                    calc = c10 + (c13 or 0)
                    if abs(calc - c14) > 0.50:
                        row_errors.append({
                            "formula": "C14 = C10 + C13",
                            "expected": calc,
                            "actual": c14,
                            "diff": abs(calc - c14)
                        })

                else:
                    print(table_idx, r_num, " НЕ ПРОВКРИЛ")
                if row_errors:
                    stats["errors_count"] += len(row_errors)
                    errors.append({
                        "table_idx": table_idx,
                        "page": page_num,
                        "row": r_num,
                        "errors": row_errors,
                        "values": {
                            "C1": c1, "C2": c2, "C3": c3, "C4": c4, "C5": c5, "C6": c6, "C7": c7,
                            "C8": c8, "C9": c9, "C10": c10, "C11": c11, "C12": c12, "C13": c13, "C14": c14
                        }
                    })
                else:
                    stats["rows_checked"] += 1

            if max_col == 10:
                #print(table_idx,r_num,max_col)
                #print(c1, c2, c3, c4, c5, c6, c7, c8, c9, c10)
                row_errors = []

                if c5 is not None and c2 is not None and c3 is not None:
                    calc = (c2 or 0) + (c3 or 0) + (c4 or 0)
                    if abs(calc - c5) > 0.50:
                        row_errors.append({
                            "formula": "C5 = C2 + C3 + C4",
                            "expected": calc,
                            "actual": c5,
                            "diff": abs(calc - c5)
                        })
                else:
                    print(table_idx, r_num, " НЕ ПРОВКРИЛ")
                if c8 is not None and c7 is not None and c6 is not None:
                    calc = (c6 or 0) - (c7 or 0)
                    if abs(calc - c8) > 0.50:
                        row_errors.append({
                            "formula": "C8 = C6 - C7",
                            "expected": calc,
                            "actual": c8,
                            "diff": abs(calc - c8)
                        })
                else:
                    print(table_idx, r_num, " НЕ ПРОВКРИЛ")
                # Формула 3: C14 = C10 + C13
                if c5 is not None and c9 is not None:
                    calc = (c5 or 0) + (c8 or 0)
                    if abs(calc - c9) > 0.50:
                        row_errors.append({
                            "formula": "C9 = C5 + C8",
                            "expected": calc,
                            "actual": c9,
                            "diff": abs(calc - c9)
                        })
                else:
                    print(table_idx, r_num, " НЕ ПРОВКРИЛ")
                if row_errors:
                    stats["errors_count"] += len(row_errors)
                    errors.append({
                        "table_idx": table_idx,
                        "page": page_num,
                        "row": r_num,
                        "errors": row_errors,
                        "values": {
                            "C1": c1, "C2": c2, "C3": c8, "C4": c9, "C5": c10, "C6": c11, "C7": c12,
                            "C8": c13, "C9": c14

                        }
                    })
                else:
                    stats["rows_checked"] += 1

            if max_col == 12:
                #print(c1, c2, c3, c4, c5, c6, c7, c8, c9, c10, c11, c12)
                #print(table_idx,r_num,max_col)

                row_errors = []

                if c11 is not None:
                    if c3 == 0 or c4 == 0 or c3 is None or c4 is None or c6 is not None:
                        if c6 is not None and c7 is not None and c8 is not None and c9 is not None:
                            calc = c6 - c7 - c8 - c9
                            if abs(calc - c11) > 0.50:
                                row_errors.append({
                                    "formula": "C11 = C6 - C7 - C8 - C9",
                                    "expected": calc,
                                    "actual": c11,
                                    "diff": abs(calc - c11)
                                })

                    else:
                        calc = c3 * c4
                        if abs(calc - c11) > 0.50:
                            row_errors.append({
                                "formula": "C11 = C3 * C4",
                                "expected": calc,
                                "actual": c11,
                                "diff": abs(calc - c11)
                            })

                else:
                    print(table_idx, r_num, " НЕ ПРОВКРИЛ")

                if c12 is not None and c11 is not None:
                    calc = c11 / c5 * c10
                    if abs(calc - c12) > 0.50:
                        row_errors.append({
                            "formula": "C11 / C5 * C10",
                            "expected": calc,
                            "actual": c12,
                            "diff": abs(calc - c12)
                        })
                else:
                    print(table_idx, r_num, " НЕ ПРОВКРИЛ")
                if row_errors:
                    stats["errors_count"] += len(row_errors)
                    errors.append({
                        "table_idx": table_idx,
                        "page": page_num,
                        "row": r_num,
                        "errors": row_errors,
                        "values": {
                            "C1": c1, "C2": c2, "C3": c8, "C4": c9, "C5": c10, "C6": c11, "C7": c12,
                            "C8": c13, "C9": c14

                        }
                    })
                else:
                    stats["rows_checked"] += 1
    return {"errors": errors, "stats": stats}


def get_entities_from_table(table: dict) -> dict:
    entities = table.get('extracted_entities', {})
    if not entities:
        entities = {
            'inn': table.get('inn'),
            'kpp': table.get('kpp'),
            'bik': table.get('bik'),
            'rs_account': table.get('rs_account'),
            'org_name': table.get('org_name')
        }
    return entities


def find_best_table_match(tables: list, qr: dict):
    if not tables:
        return None, 0, {}

    def _digits(s):
        return re.sub(r"\D+", "", s or "")

    qr_inn = _digits(qr.get('inn'))
    qr_kpp = _digits(qr.get('kpp'))
    qr_bik = _digits(qr.get('bik'))
    qr_rs = _digits(qr.get('rs_account'))

    best_table = None
    best_score = 0
    best_details = {}

    for table in tables:
        entities = get_entities_from_table(table)
        score = 0
        details = {'inn': False, 'kpp': False, 'bik': False, 'rs': False}

        tbl_inn = _digits(entities.get('inn'))
        if qr_inn and tbl_inn and qr_inn == tbl_inn:
            score += 50;
            details['inn'] = True

        tbl_kpp = _digits(entities.get('kpp'))
        if qr_kpp and tbl_kpp and qr_kpp == tbl_kpp:
            score += 30;
            details['kpp'] = True

        tbl_bik = _digits(entities.get('bik'))
        if qr_bik and tbl_bik and qr_bik == tbl_bik:
            score += 25;
            details['bik'] = True

        tbl_rs = _digits(entities.get('rs_account'))
        if qr_rs and tbl_rs and qr_rs == tbl_rs:
            score += 100;
            details['rs'] = True

        if score > best_score:
            best_score = score
            best_table = table
            best_details = details

    return best_table, best_score, best_details


def validate_qr_vs_tables(data):
    if 'ner' in data:
        tables = data['ner'].get('tables', [])
        qr_codes = data['ner'].get('qr_codes', [])
        headers = data['ner'].get('headers', [])
        print("\n📂 Формат: API")
    else:
        tables = data.get('tables', [])
        qr_codes = data.get('qr_codes', [])
        headers = data.get('headers', [])
        print("\n📂 Формат: Локальный")

    if not qr_codes or not tables:
        return [], {"matched": 0, "mismatches": 0, "no_match": 0}

    tables_by_doc = defaultdict(list)
    for table in tables:
        doc_idx = table.get("doc_idx")
        if doc_idx:
            tables_by_doc[doc_idx].append(table)

    headers_by_doc = {}
    for header in headers:
        doc_idx = header.get("doc_idx")
        if doc_idx and doc_idx not in headers_by_doc:
            headers_by_doc[doc_idx] = header

    discrepancies = []
    stats = {"matched": 0, "mismatches": 0, "no_match": 0, "header_missing_fields": 0, "header_vs_qr_mismatch": 0}

    # ✅ 1. Проверка обязательных полей в шапке
    print("НАЧИНАЮ СТРАВНИВАТЬ")
    required_header_fields = ["address", "owner", "area", "sum_owner", "ls_account"]
    for header in headers:
        doc_idx = header.get("doc_idx")
        page = header.get("page")
        entities = header.get("extracted_entities", {})
        missing = [f for f in required_header_fields if not _normalize_val(entities.get(f))]
        if missing:
            print("Проблемы с адрегом овнером ареа сумм овнер")
            discrepancies.append({
                "doc_idx": doc_idx, "page": page, "qr_idx": None,
                "type": "Отсутствуют обязательные поля в шапке",
                "missing_fields": missing,
                "header_data": {f: entities.get(f) for f in required_header_fields}
            })
            stats["header_missing_fields"] += 1

    # ✅ 2. Оригинальная логика сверки QR и таблиц (ваш код)
    print("НАЧИНАЮ СТРАВНИВАТЬ 2")
    for qr in qr_codes:
        doc_idx = qr.get("doc_idx")
        page = qr.get("page")
        qr_idx = qr.get("qr_idx")

        qr_inn_raw = qr.get('inn')
        qr_kpp_raw = qr.get('kpp')
        qr_bik_raw = qr.get('bik')
        qr_rs_raw = qr.get('rs_account')
        qr_purpose_ls = _normalize_val(qr.get('purpose'))

        qr_data_full = {
            "inn": qr_inn_raw, "kpp": qr_kpp_raw,
            "bik": qr_bik_raw, "rs": qr_rs_raw, "ls": qr_purpose_ls
        }

        # ✅ 3. Сверка Адреса и Владельца между Шапкой и QR
        print("НАЧИНАЮ СТРАВНИВАТЬ 3")
        header_obj = headers_by_doc.get(doc_idx)
        if header_obj:
            h_ent = header_obj.get("extracted_entities", {})
            h_address = _normalize_val(h_ent.get("address"))
            h_owner = _normalize_val(h_ent.get("owner"))

            qr_address = _normalize_val(qr.get("payerAddress"))
            qr_owner_parts = [
                _normalize_val(qr.get('lastName')),
                _normalize_val(qr.get('firstName')),
                _normalize_val(qr.get('middleName'))
            ]
            qr_owner = " ".join(filter(None, qr_owner_parts)).strip() or None

            # Сравнение адреса
            print("НАЧИНАЮ СТРАВНИВАТЬ АДРЕСС")
            if h_address and qr_address:
                if _normalize_for_compare(h_address) != _normalize_for_compare(qr_address):
                    discrepancies.append({
                        "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
                        "type": "Несовпадение адреса (Шапка vs QR)",
                        "header_val": h_address, "qr_val": qr_address
                    })
                    stats["header_vs_qr_mismatch"] += 1

            elif bool(h_address) != bool(qr_address):
                discrepancies.append({
                    "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
                    "type": "Адрес найден только в одном источнике",
                    "header_val": h_address, "qr_val": qr_address
                })
                stats["header_vs_qr_mismatch"] += 1

            # Сравнение владельца
            # print("НАЧИНАЮ СТРАВНИВАТЬ ВЛЕДЛЬЦА")
            # if h_owner and qr_owner:
            #     if _normalize_for_compare(h_owner) != _normalize_for_compare(qr_owner):
            #         discrepancies.append({
            #             "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
            #             "type": "Несовпадение ФИО владельца (Шапка vs QR)",
            #             "header_val": h_owner, "qr_val": qr_owner
            #         })
            #         stats["header_vs_qr_mismatch"] += 1
            #
            # elif bool(h_owner) != bool(qr_owner):
            #     discrepancies.append({
            #         "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
            #         "type": "ФИО владельца найдено только в одном источнике",
            #         "header_val": h_owner, "qr_val": qr_owner
            #     })
            #     stats["header_vs_qr_mismatch"] += 1

            print("НАЧИНАЮ СРАВНИВАТЬ ВЛАДЕЛЬЦА")
            if h_owner and qr_owner:
                # Заменена строгая проверка на устойчивую
                if not _names_match(h_owner, qr_owner):
                    discrepancies.append({
                        "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
                        "type": "Несовпадение ФИО владельца (Шапка vs QR)",
                        "header_val": h_owner, "qr_val": qr_owner
                    })
                    stats["header_vs_qr_mismatch"] += 1

            elif bool(h_owner) != bool(qr_owner):
                discrepancies.append({
                    "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
                    "type": "ФИО владельца найдено только в одном источнике",
                    "header_val": h_owner, "qr_val": qr_owner
                })
                stats["header_vs_qr_mismatch"] += 1

        # --- Ваш оригинальный код поиска таблицы и проверки реквизитов ---
        doc_tables = tables_by_doc.get(doc_idx, [])
        if not doc_tables:
            discrepancies.append({
                "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
                "type": "Нет таблиц в документе",
                "qr_data": qr_data_full, "table_data": {},
                "details": f"Не найдено таблиц для doc_idx={doc_idx}"
            })
            stats["no_match"] += 1
            continue

        target_table, score, details = find_best_table_match(doc_tables, qr)
        if not target_table or score == 0:
            discrepancies.append({
                "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
                "type": "Нет совпадений таблицы (score=0)",
                "qr_data": qr_data_full, "table_data": {},
                "details": "Таблица не найдена по реквизитам"
            })
            stats["no_match"] += 1
            continue

        table_idx = target_table.get('table_idx', '?')
        entities = get_entities_from_table(target_table)
        table_data_full = {
            "inn": entities.get('inn'), "kpp": entities.get('kpp'),
            "bik": entities.get('bik'), "rs": entities.get('rs_account')
        }

        header_ls = None
        if header_obj:
            h_entities = header_obj.get("extracted_entities", {})
            header_ls = _normalize_val(h_entities.get("ls_account"))

        ls_mismatch = False
        ls_diff_detail = None
        if header_ls and qr_purpose_ls and header_ls.lower() != qr_purpose_ls.lower():
            ls_mismatch = True
            ls_diff_detail = {"field": "Лицевой счет", "header_val": header_ls, "qr_val": qr_purpose_ls}

        diffs = []
        missing_from_table = []

        def check_field(name, qr_val, tbl_val):
            q = _normalize_val(qr_val)
            t = _normalize_val(tbl_val)
            if q and t and q != t:
                diffs.append((name, t, q))
            elif q and not t:
                missing_from_table.append(name)

        check_field('ИНН', qr_inn_raw, entities.get('inn'))
        check_field('КПП', qr_kpp_raw, entities.get('kpp'))
        check_field('БИК', qr_bik_raw, entities.get('bik'))
        check_field('Р/С', qr_rs_raw, entities.get('rs_account'))

        req_pairs = [
            ('ИНН', _normalize_val(entities.get('inn'))),
            ('КПП', _normalize_val(entities.get('kpp'))),
            ('БИК', _normalize_val(entities.get('bik'))),
            ('Р/С', _normalize_val(entities.get('rs_account'))),
        ]
        present_req = [name for name, val in req_pairs if val]
        missing_req = [name for name, val in req_pairs if not val]
        incomplete_table_requisites = bool(present_req) and bool(missing_req)

        has_errors = bool(diffs) or ls_mismatch or bool(missing_from_table) or incomplete_table_requisites
        if has_errors:
            if diffs:
                error_type = "Несовпадение реквизитов"
            elif missing_from_table or incomplete_table_requisites:
                error_type = "Неполные реквизиты в таблице"
            else:
                error_type = "Несовпадение Лицевого Счета"
            disc_item = {
                "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
                "type": error_type, "qr_data": qr_data_full, "table_data": table_data_full,
                "header_ls": header_ls, "diffs": diffs, "score": score,
                "matched_table_idx": table_idx
            }
            if missing_from_table:
                disc_item["missing_from_table"] = missing_from_table
            if incomplete_table_requisites:
                disc_item["missing_requisites"] = missing_req
            if ls_mismatch:
                disc_item['ls_error'] = ls_diff_detail
            discrepancies.append(disc_item)
            stats["mismatches"] += 1
        else:
            stats["matched"] += 1

    return discrepancies, stats


def save_report_to_docx(discrepancies: list, qr_stats: dict, formula_result: dict,
                        filename: str = "reconciliation_report.docx"):
    """Генерирует DOCX отчет с результатами сверки QR и проверки формул."""
    print("Начал отчет делать")
    try:
        doc = Document()
        title = doc.add_heading('Отчет о проверке документов', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Дата: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        doc.add_heading('1. Сверка реквизитов QR, шапки и таблиц', level=1)
        p_stats = doc.add_paragraph()
        p_stats.add_run("Статистика: ").bold = True
        p_stats.add_run(f"Всего QR: {sum(qr_stats.values())}, ")
        p_stats.add_run(f"OK: {qr_stats['matched']}, ").font.color.rgb = RGBColor(0, 128, 0)
        p_stats.add_run(f"Ошибки рек.: {qr_stats['mismatches']}").font.color.rgb = RGBColor(255, 0, 0)
        if qr_stats.get('no_match'):
            p_stats.add_run(f", Не найдено таблиц: {qr_stats['no_match']}")
        if qr_stats.get('header_vs_qr_mismatch'):
            p_stats.add_run(f", Шапка↔QR: {qr_stats['header_vs_qr_mismatch']}").font.color.rgb = RGBColor(204, 153, 0)
        if qr_stats.get('header_missing_fields'):
            p_stats.add_run(f", Пропуски шапки: {qr_stats['header_missing_fields']}").font.color.rgb = RGBColor(192, 0,
                                                                                                                0)

        if not discrepancies:
            doc.add_paragraph("✅ Ошибок сверки не найдено.").runs[0].font.color.rgb = RGBColor(0, 128, 0)
        else:
            doc.add_heading('Детализация ошибок сверки', level=2)
            for i, err in enumerate(discrepancies, 1):
                err_type = err.get("type", "Неизвестная ошибка")
                doc.add_heading(f'Ошибка #{i}: {err_type}', level=3)
                p = doc.add_paragraph()
                p.add_run(f"QR ID: #{err.get('qr_idx', 'Н/Д')} (Doc: {err.get('doc_idx')}, Page: {err.get('page')})")

                # ✅ Обработка новых типов ошибок
                if err_type == "Отсутствуют обязательные поля в шапке":
                    p = doc.add_paragraph();
                    p.add_run("❌ Отсутствующие поля:").bold = True
                    doc.add_paragraph(", ".join(err.get("missing_fields", [])))
                    continue
                if err_type in ["Несовпадение адреса (Шапка vs QR)", "Адрес найден только в одном источнике"]:
                    p = doc.add_paragraph();
                    p.add_run("📍 Адрес:").bold = True
                    doc.add_paragraph(f"Шапка: {err.get('header_val', '—')}\nQR: {err.get('qr_val', '—')}")
                    continue
                if err_type in ["Несовпадение ФИО владельца (Шапка vs QR)",
                                "ФИО владельца найдено только в одном источнике"]:
                    p = doc.add_paragraph();
                    p.add_run("👤 Владелец:").bold = True
                    doc.add_paragraph(f"Шапка: {err.get('header_val', '—')}\nQR: {err.get('qr_val', '—')}")
                    continue

                # --- Ваш оригинальный код отчёта для реквизитов ---
                p = doc.add_paragraph();
                p.add_run("📱 Данные из QR:").bold = True
                qr_d = err.get('qr_data', {})
                tbl_qr = doc.add_table(rows=1, cols=2, style='Table Grid')
                for fname, fval in [("ИНН", qr_d.get('inn')), ("КПП", qr_d.get('kpp')), ("БИК", qr_d.get('bik')),
                                    ("Р/С", qr_d.get('rs')), ("ЛС (Purpose)", qr_d.get('ls'))]:
                    row = tbl_qr.add_row().cells
                    row[0].text = fname
                    row[1].text = str(fval) if fval is not None else "—"

                if err.get('table_data'):
                    p = doc.add_paragraph();
                    p.add_run("📄 Данные из Таблицы:").bold = True
                    tbl_d = err.get('table_data', {})
                    tbl_table = doc.add_table(rows=1, cols=2, style='Table Grid')
                    for fname, fval in [("ИНН", tbl_d.get('inn')), ("КПП", tbl_d.get('kpp')),
                                        ("БИК", tbl_d.get('bik')), ("Р/С", tbl_d.get('rs'))]:
                        row = tbl_table.add_row().cells
                        row[0].text = fname
                        row[1].text = str(fval) if fval is not None else "—"

                header_ls = err.get('header_ls')
                if header_ls is not None:
                    doc.add_paragraph(f"Лицевой счет из Шапки: {header_ls}")

                for field, tbl_val, qr_val in err.get('diffs', []):
                    p_err = doc.add_paragraph(style='List Bullet')
                    p_err.add_run(f"{field}: ").bold = True
                    p_err.add_run(f"Таблица='{tbl_val}'  ≠  QR='{qr_val}'")

                if err.get('ls_error'):
                    ls_err = err['ls_error']
                    p_ls = doc.add_paragraph(style='List Bullet')
                    p_ls.add_run("Не совпадает Лицевой Счет: ").bold = True
                    p_ls.add_run(f"Шапка='{ls_err['header_val']}', QR='{ls_err['qr_val']}'")

                if err.get('details'):
                    doc.add_paragraph(f"Примечание: {err['details']}")

                doc.add_paragraph()  # отступ

        # ----- ПРОВЕРКА ФОРМУЛ В ТАБЛИЦАХ -----
        doc.add_page_break()
        doc.add_heading('2. Проверка арифметических формул в таблицах', level=1)
        fstats = formula_result['stats']
        p_f = doc.add_paragraph()
        p_f.add_run("Статистика: ").bold = True
        p_f.add_run(f"Таблиц проверено: {fstats['tables_checked']}, ")
        p_f.add_run(f"Строк без ошибок: {fstats['rows_checked']}, ")
        p_f.add_run(f"Найдено ошибок: {fstats['errors_count']}").font.color.rgb = RGBColor(255, 0, 0)

        if fstats['errors_count'] == 0:
            doc.add_paragraph("Все формулы сходятся.").runs[0].font.color.rgb = RGBColor(0, 128, 0)
        else:
            doc.add_heading('Детализация ошибок формул', level=2)
            for err in formula_result['errors']:
                doc.add_heading(f"Таблица {err['table_idx']} (стр. {err['page']}), строка {err['row']}", level=3)
                vals = err['values']
                tbl_vals = doc.add_table(rows=1, cols=2, style='Table Grid')
                for k, v in vals.items():
                    if v is not None:
                        row = tbl_vals.add_row().cells
                        row[0].text = k
                        row[1].text = f"{v:.2f}"
                for e in err['errors']:
                    p_err = doc.add_paragraph(style='List Bullet')
                    p_err.add_run(f"{e['formula']}: ").bold = True
                    p_err.add_run(
                        f"ожидалось {e['expected']:.2f}, получено {e['actual']:.2f} (разница {e['diff']:.2f})")

        doc.save(filename)
        print(f"\nОтчет сохранен: {os.path.abspath(filename)}")
        return filename
    except Exception as e:
        print(f"Ошибка создания DOCX: {e}")
        import traceback
        traceback.print_exc()
        return None

#
# def save_report_to_docx(discrepancies: list, qr_stats: dict,
#                         formula_result: dict, filename: str = "reconciliation_report.docx"):
#     """Генерирует DOCX отчет с результатами сверки QR и проверки формул."""
#     try:
#         doc = Document()
#         title = doc.add_heading('Отчет о проверке документов', 0)
#         title.alignment = WD_ALIGN_PARAGRAPH.CENTER
#
#         doc.add_paragraph(f"Дата: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
#
#         # ----- СТАТИСТИКА СВЕРКИ QR -----
#         doc.add_heading('1. Сверка реквизитов QR с таблицами', level=1)
#         p_stats = doc.add_paragraph()
#         p_stats.add_run("Статистика: ").bold = True
#         p_stats.add_run(f"Всего QR: {sum(qr_stats.values())}, ")
#         p_stats.add_run(f"OK: {qr_stats['matched']}, ").font.color.rgb = RGBColor(0,128,0)
#         p_stats.add_run(f"Ошибки: {qr_stats['mismatches']}").font.color.rgb = RGBColor(255,0,0)
#         if qr_stats.get('no_match'):
#             p_stats.add_run(f", Не найдено таблиц: {qr_stats['no_match']}")
#
#         if not discrepancies:
#             doc.add_paragraph("✅ Ошибок сверки не найдено.").runs[0].font.color.rgb = RGBColor(0,128,0)
#         else:
#             doc.add_heading('Детализация ошибок сверки', level=2)
#             for i, err in enumerate(discrepancies, 1):
#                 doc.add_heading(f'Ошибка #{i}: {err["type"]}', level=3)
#                 p = doc.add_paragraph()
#                 p.add_run(f"QR ID: #{err.get('qr_idx')} (Doc: {err.get('doc_idx')}, Page: {err.get('page')})")
#
#                 # Данные QR
#                 p = doc.add_paragraph(); p.add_run("📱 Данные из QR:").bold = True
#                 qr_d = err.get('qr_data', {})
#                 tbl_qr = doc.add_table(rows=1, cols=2, style='Table Grid')
#                 for fname, fval in [("ИНН",qr_d.get('inn')),("КПП",qr_d.get('kpp')),("БИК",qr_d.get('bik')),
#                                      ("Р/С",qr_d.get('rs')),("ЛС (Purpose)",qr_d.get('ls'))]:
#                     row = tbl_qr.add_row().cells
#                     row[0].text = fname
#                     row[1].text = str(fval) if fval is not None else "—"
#
#                 # Данные таблицы
#                 if err.get('table_data'):
#                     p = doc.add_paragraph(); p.add_run("📄 Данные из Таблицы:").bold = True
#                     tbl_d = err.get('table_data', {})
#                     tbl_table = doc.add_table(rows=1, cols=2, style='Table Grid')
#                     for fname, fval in [("ИНН",tbl_d.get('inn')),("КПП",tbl_d.get('kpp')),
#                                          ("БИК",tbl_d.get('bik')),("Р/С",tbl_d.get('rs'))]:
#                         row = tbl_table.add_row().cells
#                         row[0].text = fname
#                         row[1].text = str(fval) if fval is not None else "—"
#
#                 # Лицевой счёт из шапки
#                 header_ls = err.get('header_ls')
#                 if header_ls is not None:
#                     doc.add_paragraph(f"Лицевой счет из Шапки: {header_ls}")
#
#                 # Конкретные несовпадения
#                 for field, tbl_val, qr_val in err.get('diffs', []):
#                     p_err = doc.add_paragraph(style='List Bullet')
#                     p_err.add_run(f"{field}: ").bold = True
#                     p_err.add_run(f"Таблица='{tbl_val}'  ≠  QR='{qr_val}'")
#
#                 if err.get('ls_error'):
#                     ls_err = err['ls_error']
#                     p_ls = doc.add_paragraph(style='List Bullet')
#                     p_ls.add_run("Не совпадает Лицевой Счет: ").bold = True
#                     p_ls.add_run(f"Шапка='{ls_err['header_val']}', QR='{ls_err['qr_val']}'")
#
#                 if err.get('details'):
#                     doc.add_paragraph(f"Примечание: {err['details']}")
#
#                 doc.add_paragraph()  # отступ
#
#         # ----- ПРОВЕРКА ФОРМУЛ В ТАБЛИЦАХ -----
#         doc.add_page_break()
#         doc.add_heading('2. Проверка арифметических формул в таблицах', level=1)
#         fstats = formula_result['stats']
#         p_f = doc.add_paragraph()
#         p_f.add_run("Статистика: ").bold = True
#         p_f.add_run(f"Таблиц проверено: {fstats['tables_checked']}, ")
#         p_f.add_run(f"Строк без ошибок: {fstats['rows_checked']}, ")
#         p_f.add_run(f"Найдено ошибок: {fstats['errors_count']}").font.color.rgb = RGBColor(255,0,0)
#
#         if fstats['errors_count'] == 0:
#             doc.add_paragraph("Все формулы сходятся.").runs[0].font.color.rgb = RGBColor(0,128,0)
#         else:
#             doc.add_heading('Детализация ошибок формул', level=2)
#             for err in formula_result['errors']:
#                 doc.add_heading(f"Таблица {err['table_idx']} (стр. {err['page']}), строка {err['row']}", level=3)
#                 # Покажем значения ячеек
#                 vals = err['values']
#                 tbl_vals = doc.add_table(rows=1, cols=2, style='Table Grid')
#                 for k,v in vals.items():
#                     if v is not None:
#                         row = tbl_vals.add_row().cells
#                         row[0].text = k
#                         row[1].text = f"{v:.2f}"
#                 # Ошибки
#                 for e in err['errors']:
#                     p_err = doc.add_paragraph(style='List Bullet')
#                     p_err.add_run(f"{e['formula']}: ").bold = True
#                     p_err.add_run(f"ожидалось {e['expected']:.2f}, получено {e['actual']:.2f} (разница {e['diff']:.2f})")
#
#         doc.save(filename)
#         print(f"\nОтчет сохранен: {os.path.abspath(filename)}")
#         return filename
#     except Exception as e:
#         print(f"Ошибка создания DOCX: {e}")
#         import traceback
#         traceback.print_exc()
#         return None

# ----------------------------------------------------------------------
# 5. ГЛАВНАЯ ФУНКЦИЯ (объединяет всё)
# ----------------------------------------------------------------------
# def main(json_path: str, check_qr: bool = True, check_formulas: bool = True):
#     """Основная точка входа."""
#     if not os.path.exists(json_path):
#         print(f"❌ Файл {json_path} не найден!")
#         return
#
#     with open(json_path, 'r', encoding='utf-8') as f:
#         data = json.load(f)
#
#     discrepancies = []
#     qr_stats = {"matched": 0, "mismatches": 0, "no_match": 0}
#     formula_result = {"errors": [], "stats": {"tables_checked": 0, "rows_checked": 0, "errors_count": 0}}
#
#     if check_qr:
#         print("\n🔍 ЗАПУСК СВЕРКИ QR С ТАБЛИЦАМИ")
#         discrepancies, qr_stats = validate_qr_vs_tables(data)
#     else:
#         print("\n⏩ Сверка QR пропущена")
#
#     if check_formulas:
#         print("\n🔍 ЗАПУСК ПРОВЕРКИ ФОРМУЛ В ТАБЛИЦАХ")
#         # Извлекаем список таблиц (учитываем возможную вложенность в 'ner')
#         if 'ner' in data:
#             tables = data['ner'].get('tables', [])
#         else:
#             tables = data.get('tables', [])
#         formula_result = check_table_formulas(tables)
#         fs = formula_result['stats']
#         print(f"   Таблиц проверено: {fs['tables_checked']}")
#         print(f"   Строк без ошибок: {fs['rows_checked']}")
#         print(f"   Найдено ошибок: {fs['errors_count']}")
#     else:
#         print("\n⏩ Проверка формул пропущена")
#
#     # Сохраняем отчёт
#     save_report_to_docx(discrepancies, qr_stats, formula_result,
#                         filename="full_verification_report.docx")
#
# if __name__ == "__main__":
#     # Пример использования: можно указать путь к JSON и флаги
#     JSON_FILE = "taroe/ner_results_spacy.json"   # измените на свой путь
#     # main(JSON_FILE, check_qr=True, check_formulas=True)
#
#     # Для теста с аргументами командной строки
#     if len(sys.argv) > 1:
#         JSON_FILE = sys.argv[1]
#     main(JSON_FILE, check_qr=True, check_formulas=True)
