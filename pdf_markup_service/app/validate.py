import json
import os
import re
import sys
import datetime
from collections import defaultdict

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Требуется библиотека python-docx. Установите: pip install python-docx")
    sys.exit(1)

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")


def _normalize_val(val):
    """Приводит значение к строке, убирает пробелы. Если пусто — возвращает None для единообразия"""
    if val is None:
        return None
    s = str(val).strip()
    return s if s else None


def get_entities_from_table(table: dict) -> dict:
    """Извлекает реквизиты из таблицы"""
    entities = table.get('extracted_entities', {})
    if not entities:
        entities = {
            'inn': table.get('inn'),
            'kpp': table.get('kpp'),
            'bik': table.get('bik'),
            'rs_account': table.get('rs_account'),
            'org_name': table.get('org_name')
        }
    # Возвращаем как есть, фильтрацию сделаем при формировании отчета, чтобы видеть отсутствующие
    return entities


def find_best_table_match(tables: list, qr: dict):
    """Умное сопоставление QR с таблицей через scoring."""
    if not tables:
        return None, 0, {}

    # Для скоринга используем только цифры
    def _digits(s):
        return re.sub(r"\D+", "", s or "")

    qr_inn = _digits(qr.get('inn'))
    qr_kpp = _digits(qr.get('kpp'))
    qr_bik = _digits(qr.get('bik'))
    qr_rs = _digits(qr.get('rs_account'))

    best_table = None
    best_score = 0
    best_details = {}

    for table in tables:
        entities = get_entities_from_table(table)
        score = 0
        details = {'inn': False, 'kpp': False, 'bik': False, 'rs': False}

        tbl_inn = _digits(entities.get('inn'))
        if qr_inn and tbl_inn and qr_inn == tbl_inn:
            score += 50;
            details['inn'] = True

        tbl_kpp = _digits(entities.get('kpp'))
        if qr_kpp and tbl_kpp and qr_kpp == tbl_kpp:
            score += 30;
            details['kpp'] = True

        tbl_bik = _digits(entities.get('bik'))
        if qr_bik and tbl_bik and qr_bik == tbl_bik:
            score += 25;
            details['bik'] = True

        tbl_rs = _digits(entities.get('rs_account'))
        if qr_rs and tbl_rs and qr_rs == tbl_rs:
            score += 100;
            details['rs'] = True

        if score > best_score:
            best_score = score
            best_table = table
            best_details = details

    return best_table, best_score, best_details


def validate_qr_vs_tables(data: dict):
    """
    Сравнивает реквизиты из QR с таблицами и заголовками (ЛС).
    """
    # 1. Извлечение данных
    if 'ner' in data:
        tables = data['ner'].get('tables', [])
        qr_codes = data['ner'].get('qr_codes', [])
        headers = data['ner'].get('headers', [])
        print("\n📂 Формат: API")
    else:
        tables = data.get('tables', [])
        qr_codes = data.get('qr_codes', [])
        headers = data.get('headers', [])
        print("\n📂 Формат: Локальный")

    if not qr_codes:
        return [], {"matched": 0, "mismatches": 0, "no_match": 0}

    if not tables:
        return [], {"matched": 0, "mismatches": 0, "no_match": 0}

    # 2. Индексация
    tables_by_doc = defaultdict(list)
    for table in tables:
        doc_idx = table.get("doc_idx")
        if doc_idx:
            tables_by_doc[doc_idx].append(table)

    headers_by_doc = {}
    for header in headers:
        doc_idx = header.get("doc_idx")
        if doc_idx and doc_idx not in headers_by_doc:
            headers_by_doc[doc_idx] = header

    discrepancies = []
    stats = {"matched": 0, "mismatches": 0, "no_match": 0}

    for qr in qr_codes:
        doc_idx = qr.get("doc_idx")
        page = qr.get("page")
        qr_idx = qr.get("qr_idx")

        # --- Сбор данных QR (сохраняем всё, даже пустое) ---
        qr_inn_raw = qr.get('inn')
        qr_kpp_raw = qr.get('kpp')
        qr_bik_raw = qr.get('bik')
        qr_rs_raw = qr.get('rs_account')
        qr_purpose_ls = _normalize_val(qr.get('purpose'))  # ЛС из purpose

        qr_data_full = {
            "inn": qr_inn_raw,
            "kpp": qr_kpp_raw,
            "bik": qr_bik_raw,
            "rs": qr_rs_raw,
            "ls": qr_purpose_ls
        }

        # --- ПРОВЕРКА 1: Наличие таблицы ---
        doc_tables = tables_by_doc.get(doc_idx, [])
        if not doc_tables:
            discrepancies.append({
                "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
                "type": "Нет таблиц в документе",
                "qr_data": qr_data_full,
                "table_data": {},  # Пустая таблица
                "details": f"Не найдено таблиц для doc_idx={doc_idx}"
            })
            stats["no_match"] += 1
            continue

        # --- ПРОВЕРКА 2: Поиск лучшей таблицы ---
        target_table, score, details = find_best_table_match(doc_tables, qr)

        if not target_table or score == 0:
            discrepancies.append({
                "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
                "type": "Нет совпадений таблицы (score=0)",
                "qr_data": qr_data_full,
                "table_data": {},
                "details": "Таблица не найдена по реквизитам"
            })
            stats["no_match"] += 1
            continue

        table_idx = target_table.get('table_idx', '?')
        entities = get_entities_from_table(target_table)

        # --- Сбор данных Таблицы (сохраняем всё, даже пустое) ---
        tbl_inn_raw = entities.get('inn')
        tbl_kpp_raw = entities.get('kpp')
        tbl_bik_raw = entities.get('bik')
        tbl_rs_raw = entities.get('rs_account')

        table_data_full = {
            "inn": tbl_inn_raw,
            "kpp": tbl_kpp_raw,
            "bik": tbl_bik_raw,
            "rs": tbl_rs_raw
        }

        # --- ПРОВЕРКА 3: Лицевой счет (Header vs QR Purpose) ---
        header_ls = None
        header_obj = headers_by_doc.get(doc_idx)
        if header_obj:
            h_entities = header_obj.get("extracted_entities", {})
            header_ls = _normalize_val(h_entities.get("ls_account"))

        ls_mismatch = False
        ls_diff_detail = None

        # Сравниваем, если оба значения присутствуют
        if header_ls and qr_purpose_ls:
            if header_ls.lower() != qr_purpose_ls.lower():
                ls_mismatch = True
                ls_diff_detail = {
                    "field": "Лицевой счет",
                    "header_val": header_ls,
                    "qr_val": qr_purpose_ls
                }

        # --- ПРОВЕРКА 4: Основные реквизиты (ИНН, КПП, БИК, РС) ---
        diffs = []

        # Helper для сравнения
        def check_field(name, qr_val, tbl_val):
            q = _normalize_val(qr_val)
            t = _normalize_val(tbl_val)
            # Если оба есть и не равны -> ошибка
            if q and t and q != t:
                diffs.append((name, t, q))
            # Если одного нет, это не добавляем в diffs (так как это не несоответствие, а отсутствие),
            # но в отчете это будет видно благодаря полному выводу данных.

        check_field('ИНН', qr_inn_raw, tbl_inn_raw)
        check_field('КПП', qr_kpp_raw, tbl_kpp_raw)
        check_field('БИК', qr_bik_raw, tbl_bik_raw)
        check_field('Р/С', qr_rs_raw, tbl_rs_raw)

        # --- ИТОГ ---
        has_errors = bool(diffs) or ls_mismatch

        if has_errors:
            error_type = "Несовпадение реквизитов"
            if ls_mismatch and not diffs:
                error_type = "Несовпадение Лицевого Счета"

            disc_item = {
                "qr_idx": qr_idx, "doc_idx": doc_idx, "page": page,
                "type": error_type,
                "qr_data": qr_data_full,  # Полный набор данных QR
                "table_data": table_data_full,  # Полный набор данных Таблицы
                "header_ls": header_ls,  # ЛС из шапки отдельно
                "diffs": diffs,
                "score": score,
                "matched_table_idx": table_idx
            }

            if ls_mismatch:
                disc_item['ls_error'] = ls_diff_detail

            discrepancies.append(disc_item)
            stats["mismatches"] += 1
        else:
            stats["matched"] += 1

    print("\n" + "-" * 70)
    print(f"🧾 ИТОГИ:")
    print(f"   ✅ Совпадений: {stats['matched']}")
    print(f"   ❌ Несовпадений: {stats['mismatches']}")
    print("=" * 70)

    return discrepancies, stats


def save_report_to_docx(discrepancies: list, stats: dict, filename: str = "reconciliation_report.docx"):
    """Генерирует DOCX отчет с ПОЛНЫМ перечнем данных"""
    try:
        doc = Document()
        title = doc.add_heading('Отчет о сверке реквизитов', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Дата: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        p_stats = doc.add_paragraph()
        p_stats.add_run("Статистика: ").bold = True
        p_stats.add_run(f"Всего: {sum(stats.values())}, ")
        p_stats.add_run(f"OK: {stats['matched']}, ").font.color.rgb = RGBColor(0, 128, 0)
        p_stats.add_run(f"Errors: {stats['mismatches']}").font.color.rgb = RGBColor(255, 0, 0)

        doc.add_paragraph("_" * 30)

        if not discrepancies:
            doc.add_paragraph("✅ Ошибок не найдено.").runs[0].font.color.rgb = RGBColor(0, 128, 0)
        else:
            doc.add_heading('Детализация ошибок', level=1)

            for i, err in enumerate(discrepancies, 1):
                h = doc.add_heading(f'Ошибка #{i}: {err["type"]}', level=2)

                p = doc.add_paragraph()
                p.add_run("QR ID: ").bold = True
                p.add_run(f"#{err.get('qr_idx')} (Doc: {err.get('doc_idx')}, Page: {err.get('page')})")

                # --- БЛОК ДАННЫХ QR (ВСЕ ПОЛЯ) ---
                p = doc.add_paragraph()
                run_qr = p.add_run("📱 Данные из QR:")
                run_qr.bold = True
                run_qr.font.color.rgb = RGBColor(0, 0, 255)  # Синий для QR

                qr_d = err.get('qr_data', {})
                fields = [
                    ("ИНН", qr_d.get('inn')),
                    ("КПП", qr_d.get('kpp')),
                    ("БИК", qr_d.get('bik')),
                    ("Р/С", qr_d.get('rs')),
                    ("ЛС (Purpose)", qr_d.get('ls'))
                ]

                tbl_qr = doc.add_table(rows=1, cols=2, style='Table Grid')
                tbl_qr.autofit = True
                for fname, fval in fields:
                    row = tbl_qr.add_row().cells
                    row[0].text = fname
                    row[1].text = str(fval) if fval is not None else "—"

                # --- БЛОК ДАННЫХ ТАБЛИЦЫ (ВСЕ ПОЛЯ) ---
                if err.get('table_data'):
                    p = doc.add_paragraph()
                    run_tbl = p.add_run("📄 Данные из Таблицы:")
                    run_tbl.bold = True
                    run_tbl.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый для Таблицы

                    tbl_d = err.get('table_data', {})
                    fields_tbl = [
                        ("ИНН", tbl_d.get('inn')),
                        ("КПП", tbl_d.get('kpp')),
                        ("БИК", tbl_d.get('bik')),
                        ("Р/С", tbl_d.get('rs'))
                    ]

                    tbl_table = doc.add_table(rows=1, cols=2, style='Table Grid')
                    tbl_table.autofit = True
                    for fname, fval in fields_tbl:
                        row = tbl_table.add_row().cells
                        row[0].text = fname
                        row[1].text = str(fval) if fval is not None else "—"

                # --- БЛОК ЛИЦЕВОГО СЧЕТА ИЗ ШАПКИ ---
                header_ls = err.get('header_ls')
                if header_ls is not None or err.get('ls_error'):
                    p = doc.add_paragraph()
                    run_hdr = p.add_run("📑 Лицевой счет из Шапки (Header):")
                    run_hdr.bold = True
                    run_hdr.font.color.rgb = RGBColor(128, 0, 128)  # Фиолетовый

                    p_ls = doc.add_paragraph()
                    p_ls.add_run(f"Значение: {header_ls if header_ls else '—'}")

                # --- ВЫВОД КОНКРЕТНЫХ ОШИБОК ---
                if err.get('diffs'):
                    p = doc.add_paragraph()
                    run_err = p.add_run("❌ Выявленные несовпадения реквизитов:")
                    run_err.bold = True
                    run_err.font.color.rgb = RGBColor(255, 0, 0)

                    for field, tbl_val, qr_val in err['diffs']:
                        p_err = doc.add_paragraph(style='List Bullet')
                        p_err.add_run(f"{field}: ").bold = True
                        p_err.add_run(f"Таблица='{tbl_val}'  ≠  QR='{qr_val}'")

                if err.get('ls_error'):
                    ls_err = err['ls_error']
                    p = doc.add_paragraph()
                    run_ls = p.add_run("❌ Не совпадает Лицевой Счет:")
                    run_ls.bold = True
                    run_ls.font.color.rgb = RGBColor(255, 0, 0)

                    p_ls_err = doc.add_paragraph(style='List Bullet')
                    p_ls_err.add_run(f"В Шапке: '{ls_err['header_val']}'\n")
                    p_ls_err.add_run(f"В QR (Purpose): '{ls_err['qr_val']}'")

                if err.get('details'):
                    p = doc.add_paragraph()
                    p.add_run("💡 Примечание: ").bold = True
                    p.add_run(err['details'])

                doc.add_paragraph()  # Отступ между ошибками
                doc.add_page_break()  # Каждая ошибка с новой страницы для читаемости (опционально)

        file_path = os.path.join(os.getcwd(), filename)
        doc.save(file_path)
        print(f"📄 Отчет сохранен: {file_path}")
        return file_path

    except Exception as e:
        print(f"❌ Ошибка создания DOCX: {e}")
        import traceback
        traceback.print_exc()
        return None